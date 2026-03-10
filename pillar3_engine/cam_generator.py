"""
Credit Appraisal Memo (CAM) Generator for Intelli-Credit.
Generates a professional Word document with 5 Cs analysis, SHAP chart,
financial highlights, research findings, and final decision.
"""
import os
from datetime import datetime
from pathlib import Path
from config import OUTPUT_DIR, IS_DEMO, has_openai_key


def generate_cam(company_data: dict, five_cs: dict, scoring: dict,
                 research: dict, gst_validation: dict = None,
                 qualitative_notes: list = None,
                 swot_analysis: str = None,
                 committee_verdict: dict = None) -> str:
    """
    Generate a Credit Appraisal Memo as a Word document.

    Returns: path to the generated .docx file
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        return _generate_text_cam(company_data, five_cs, scoring, research, gst_validation, qualitative_notes, swot_analysis)

    doc = Document()

    # --- Styles ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)

    # --- Title Page ---
    title = doc.add_heading("CREDIT APPRAISAL MEMORANDUM", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(company_data.get("company_name", "Company Name"))
    run.font.size = Pt(18)
    run.bold = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"CIN: {company_data.get('cin', 'N/A')}").font.size = Pt(11)
    meta.add_run(f"\nGenerated: {datetime.now().strftime('%d-%b-%Y %H:%M')}").font.size = Pt(10)
    meta.add_run(f"\nPowered by Intelli-Credit AI Engine").font.size = Pt(9)

    doc.add_page_break()

    # --- Executive Summary ---
    doc.add_heading("1. Executive Summary", level=1)

    decision = scoring.get("decision", "REFERRED")
    score = scoring.get("credit_score", 0)
    loan = company_data.get("loan_request", {})

    summary_text = (
        f"Decision: {decision}\n"
        f"Credit Score: {score:.0f}/100 (Threshold: 60)\n"
        f"Loan Request: {loan.get('type', 'Term Loan')} — ₹ {loan.get('amount_cr', 0)} Crore "
        f"@ {loan.get('proposed_rate_pct', 0)}% p.a. ({loan.get('tenor_years', 0)}-year tenor)\n"
        f"Purpose: {loan.get('purpose', 'N/A')}"
    )
    doc.add_paragraph(summary_text)

    # Decision box
    decision_para = doc.add_paragraph()
    run = decision_para.add_run(f"  RECOMMENDATION: {decision}  ")
    run.bold = True
    run.font.size = Pt(14)
    if decision == "APPROVED":
        run.font.color.rgb = RGBColor(0x28, 0xa7, 0x45)
    elif decision == "REJECTED":
        run.font.color.rgb = RGBColor(0xdc, 0x35, 0x45)
    else:
        run.font.color.rgb = RGBColor(0xff, 0xc1, 0x07)

    # --- 5 Cs Analysis ---
    doc.add_heading("2. Five Cs of Credit Analysis", level=1)

    scores = five_cs.get("scores", {})
    features = five_cs.get("features", {})

    for c_name, c_label in [
        ("character", "Character — Trustworthiness & Track Record"),
        ("capacity", "Capacity — Ability to Repay"),
        ("capital", "Capital — Financial Strength"),
        ("collateral", "Collateral — Security Coverage"),
        ("conditions", "Conditions — External Environment"),
    ]:
        doc.add_heading(c_label, level=2)
        c_score = scores.get(c_name, 0)
        doc.add_paragraph(f"Score: {c_score:.0f}/100")

        # Add features as table
        c_features = features.get(c_name, {})
        if c_features:
            table = doc.add_table(rows=1, cols=2)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "Metric"
            hdr[1].text = "Value"

            for key, val in c_features.items():
                row = table.add_row().cells
                row[0].text = key.replace("_", " ").title()
                row[1].text = str(val)

    # --- Financial Highlights ---
    doc.add_heading("3. Financial Highlights", level=1)

    financials = company_data.get("financials", {})
    if financials:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Metric"
        hdr[1].text = "FY 2022"
        hdr[2].text = "FY 2023"
        hdr[3].text = "FY 2024"

        metrics = ["revenue_cr", "ebitda_cr", "pat_cr", "total_debt_cr", "net_worth_cr",
                    "dscr", "icr", "de_ratio", "current_ratio"]

        for metric in metrics:
            row = table.add_row().cells
            row[0].text = metric.replace("_cr", " (₹ Cr)").replace("_pct", " %").replace("_", " ").title()
            for i, fy in enumerate(["fy_2022", "fy_2023", "fy_2024"]):
                val = financials.get(fy, {}).get(metric, "N/A")
                if isinstance(val, dict):
                    val = val.get("value", "N/A")
                row[i + 1].text = f"{val:.2f}" if isinstance(val, (int, float)) else str(val)

    # --- SHAP Analysis ---
    doc.add_heading("4. Scoring Model — Feature Contributions", level=1)

    doc.add_paragraph(scoring.get("explanation", "No explanation available."))

    # Embed SHAP chart if available
    chart_path = scoring.get("shap_chart_path", "")
    if chart_path and os.path.exists(chart_path):
        doc.add_picture(chart_path, width=Inches(6))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Research Findings ---
    doc.add_heading("5. Research Findings", level=1)

    # News
    news = research.get("news", {})
    doc.add_heading("5.1 News & Sentiment Analysis", level=2)
    doc.add_paragraph(news.get("summary", "No news data available."))

    articles = news.get("articles", [])
    if articles:
        for article in articles[:5]:
            doc.add_paragraph(
                f"• [{article.get('sentiment', 'neutral').upper()}] {article.get('title', '')} "
                f"— {article.get('source', '')}",
                style="List Bullet"
            )

    # MCA
    mca = research.get("mca", {})
    doc.add_heading("5.2 MCA21 Compliance", level=2)
    doc.add_paragraph(mca.get("summary", "No MCA data available."))

    # Litigation
    litigation = research.get("litigation", {})
    doc.add_heading("5.3 Litigation Analysis", level=2)
    doc.add_paragraph(litigation.get("summary", "No litigation data available."))

    # Sector
    sector = research.get("sector", {})
    doc.add_heading("5.4 Industry & Regulatory Analysis", level=2)
    doc.add_paragraph(sector.get("summary", "No sector data available."))

    # --- GST Compliance ---
    if gst_validation:
        doc.add_heading("6. GST Compliance Analysis", level=1)
        summary = gst_validation.get("summary", {})
        doc.add_paragraph(
            f"GSTIN: {gst_validation.get('gstin', 'N/A')}\n"
            f"GSTR-3B Turnover: ₹ {summary.get('gstr_3b_turnover_cr', 0)} Cr\n"
            f"Bank Credits: ₹ {summary.get('bank_credit_entries_cr', 0)} Cr\n"
            f"Turnover Variance: {summary.get('turnover_variance_pct', 0):.1f}%\n"
            f"Compliance Score: {gst_validation.get('compliance_score', 0):.0f}/100\n"
            f"Risk Level: {gst_validation.get('risk_level', 'N/A')}"
        )

        flags = gst_validation.get("flags", [])
        for flag in flags:
            doc.add_paragraph(f"⚠ {flag.get('message', '')}", style="List Bullet")

    # --- AI Credit Committee Deliberation ---
    if committee_verdict:
        doc.add_heading("7. AI Credit Committee Deliberation", level=1)
        doc.add_paragraph(f"Decision: {committee_verdict.get('verdict', 'REFER')}")
        doc.add_paragraph(f"Confidence: {committee_verdict.get('confidence', 0.8)*100:.0f}%")
        
        doc.add_heading("Triangulation Check", level=2)
        doc.add_paragraph(committee_verdict.get("triangulation_check", "N/A"))
        
        doc.add_heading("Investment Rationale", level=2)
        doc.add_paragraph(committee_verdict.get("rationale", "N/A"))
        
        doc.add_heading("Strengths & Concerns", level=2)
        for s in committee_verdict.get("key_strengths", []):
            doc.add_paragraph(f"✅ {s}", style="List Bullet")
        for c in committee_verdict.get("key_concerns", []):
            doc.add_paragraph(f"❌ {c}", style="List Bullet")

    # --- Integrated SWOT Analysis ---
    if swot_analysis:
        doc.add_heading("8. Integrated SWOT Analysis", level=1)
        for line in swot_analysis.split('\n'):
            line = line.strip()
            if not line:
                continue
            if line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                r = p.add_run(line.strip('*'))
                r.bold = True
            elif line.startswith('•') or line.startswith('-'):
                doc.add_paragraph(line.lstrip('•- '), style="List Bullet")
            else:
                doc.add_paragraph(line)

    # --- Qualitative Notes ---
    if qualitative_notes:
        doc.add_heading("9. Credit Officer's Qualitative Notes", level=1)
        for note in qualitative_notes:
            doc.add_paragraph(f"• {note}", style="List Bullet")

    # --- Research Provenance ---
    doc.add_heading("10. Decision Audit Trail", level=1)
    doc.add_paragraph("The following data sources were consulted during this credit appraisal:")

    provenance = research.get("provenance", [])
    if provenance:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Agent"
        hdr[1].text = "Method"
        hdr[2].text = "Risk Score"
        hdr[3].text = "Summary"

        for p in provenance:
            row = table.add_row().cells
            row[0].text = p.get("agent", "")
            row[1].text = p.get("method", "")
            row[2].text = str(p.get("risk_score", ""))
            row[3].text = p.get("summary", "")[:100]

    # --- Footer ---
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("— End of Credit Appraisal Memorandum —")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    run2 = footer.add_run("\nGenerated by Intelli-Credit AI Engine | Confidential")
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)

    # Save
    company_slug = company_data.get("company_name", "company").replace(" ", "_")[:30]
    filename = f"CAM_{company_slug}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    output_path = str(OUTPUT_DIR / filename)
    doc.save(output_path)

    return output_path


def _generate_text_cam(company_data, five_cs, scoring, research, gst_validation, qualitative_notes, swot_analysis) -> str:
    """Fallback: generate plain text CAM if python-docx is not available."""
    lines = [
        "=" * 60,
        "CREDIT APPRAISAL MEMORANDUM",
        "=" * 60,
        f"Company: {company_data.get('company_name', 'N/A')}",
        f"CIN: {company_data.get('cin', 'N/A')}",
        f"Decision: {scoring.get('decision', 'N/A')}",
        f"Credit Score: {scoring.get('credit_score', 0):.0f}/100",
        "",
        scoring.get("explanation", ""),
        "",
        "--- SWOT Analysis ---",
        swot_analysis or "No SWOT analysis available.",
        "",
        "=" * 60,
    ]

    output_path = str(OUTPUT_DIR / "CAM_report.txt")
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    return output_path
